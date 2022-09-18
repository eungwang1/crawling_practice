
import requests
from bs4 import BeautifulSoup
import time
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl import Workbook
from docx import Document
import time
from bs4 import BeautifulSoup
import requests


def lastPageCheck(soup: BeautifulSoup):
    isLastPage = soup.select_one("a.btn_next").attrs['aria-disabled']
    if isLastPage == 'true':
        print("마지막 페이지 입니다.")
        return True
    return False


def naver_news_crawling(url: str, document: Document, excel: Workbook, row: int):
    response = requests.get(
        url)
    html = response.text
    soup_page = BeautifulSoup(html, 'html.parser')
    articles = soup_page.select("div.info_group")

    for article in articles:
        links = article.select("a.info")
        if len(links) >= 2:  # 링크가 2개 이상이면 네이버 뉴스
            url = links[1].attrs['href']
            response = requests.get(url, headers={'User-agent': 'Mozila/5.0'})
            html = response.text
            soup_detail = BeautifulSoup(html, "html.parser")
            if "entertain" in response.url:
                title = soup_detail.select_one(".end_tit")
                content = soup_detail.select_one("#articeBody")
            elif "sports" in response.url:
                title = soup_detail.select_one("h4.title")
                content = soup_detail.select_one("#newsEndContents")
                # 불필요한 div, p 삭제
                divTags = content.select("div")
                pTags = content.select("p")
                for div in divTags:
                    div.decompose()
                for p in pTags:
                    p.decompose()
            else:
                title = soup_detail.select_one(".media_end_head_title")
                content = soup_detail.select_one("#newsct_article")
            print('============== 제목 =============')
            print(title.text.strip())
            print('============== url =============')
            print(url)
            print('============== 내용 =============')
            print(content.text.strip())
            if document:
                document.add_heading(title.text.strip(), level=0)
                document.add_paragraph(url)
                document.add_paragraph(content.text.strip())
                document.add_page_break()
            if excel:
                row += 1
                excel[f"A{str(row)}"] = title.text.strip()
                excel[f"B{str(row)}"] = content.text.strip()
                excel[f"B{str(row)}"].alignment = Alignment(
                    wrap_text=True)  # 자동줄바꿈
                excel[f"C{str(row)}"] = url
            time.sleep(0.3)
    return [row, soup_page]


def crawlingSaveAsDoc():
    try:
        keyword = input_keyword.get()
        lastPage = int(input_page.get())
        document = Document()
        # 0 - 1페이지, 11- 2페이지
        for i in range(lastPage):
            itemCount = str(i*10 + 1)
            url = f"https://search.naver.com/search.naver?sm=tab_hty.top&where=news&query={keyword}&start={itemCount}"
            row, soup_page = naver_news_crawling(url, document, None, None)
            if lastPageCheck(soup_page) == True:
                break
        common.util.createDirectory('news')
        document.save(f"./news/{keyword}.docx")
    except ValueError as e:
        print(e)
        print('입력값을 확인하세요')


def crawlingSaveAsExcel():
    keyword = input_keyword.get()
    lastPage = int(input_page.get())
    row = 1
    try:
        wb = Workbook()
        ws = wb.create_sheet("news")
        ws.column_dimensions['A'].width = 60
        ws.column_dimensions['B'].width = 60
        ws.column_dimensions['C'].width = 60
        ws['A1'] = "title"
        ws['B1'] = "content"
        ws['C1'] = "url"
        # 0 - 1페이지, 11- 2페이지
        for i in range(lastPage):
            itemCount = str(i*10 + 1)
            url = f"https://search.naver.com/search.naver?sm=tab_hty.top&where=news&query={keyword}&start={itemCount}"
            row, soup_page = naver_news_crawling(url, None, ws, row)
            if lastPageCheck(soup_page) == True:
                break
        common.util.createDirectory('news')
        wb.save(f"./news/{keyword}.xlsx")
    except ValueError as e:
        print(e)
        print('입력값을 확인하세요')


root = Tk()
root.geometry("800x800")
frm = ttk.Frame(root, padding=10)

# page 라벨
label_page = Label(root)
label_page.config(text="page")
label_page.pack()


# page 입력창
input_page = Entry(root)
input_page.pack()

# keyword 라벨
label_keyword = Label(root)
label_keyword.config(text="keyword")
label_keyword.pack()

# keyword 입력창
input_keyword = Entry(root)
input_keyword.pack()


# 실행 버튼
btn1 = Button(root)
btn1.config(text="crawling-word-save")
btn1.config(command=crawlingSaveAsDoc)
btn1.pack()

# 실행 버튼
btn1 = Button(root)
btn1.config(text="crawling-excel-save")
btn1.config(command=crawlingSaveAsExcel)
btn1.pack()

root.mainloop()
